from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "renders/house-of-wander-team-update/20260523-073132-v01/house-of-wander-team-update.docx"


def set_font(run, name="Arial", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_border(cell, color="DADCE0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = f"w:{m}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (20, "000000", 20, 6),
        "Heading 2": (16, "000000", 18, 6),
        "Heading 3": (14, "434343", 16, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("House of Wander / Casa Cabane - Team Continuation Update")
    set_font(run, size=26, color="000000", bold=False)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    for label, value in (
        ("Prepared for: ", "Joey & Jo; House of Wander team: Maaike & Laurens; Artificial Monks support"),
        ("Date: ", "May 23, 2026"),
        ("Status: ", "Local guided prototype built, higher-quality video installed, and ready for team review"),
    ):
        label_run = meta.add_run(label)
        set_font(label_run, bold=True)
        value_run = meta.add_run(value)
        set_font(value_run)
        if label != "Status: ":
            sep = meta.add_run(" | ")
            set_font(sep, color="555555")


def para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)
    return p


def bullet(doc, text, bold_label=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_label:
        r = p.add_run(bold_label)
        set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)
    return p


def numbered(doc, text, bold_label=None):
    p = doc.add_paragraph(style="List Number")
    if bold_label:
        r = p.add_run(bold_label)
        set_font(r, bold=True)
    r = p.add_run(text)
    set_font(r)
    return p


def code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, name="Courier New", size=9.5, color="434343")
    return p


def simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_font(r, size=10, bold=True)
        set_cell_border(cell)
        set_cell_margins(cell)
        shade_cell(cell, "F8F9FA")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(value)
            set_font(r, size=10)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build():
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("Executive Summary", level=1)
    para(
        doc,
        "We now have a focused local prototype for Casa Cabane inside the Dropbox project folder. The prototype turns Maaike & Laudi's real Casa Cabane walkthrough into a guided, full-screen web experience instead of a normal scroll page or Airbnb-style listing.",
    )
    para(
        doc,
        "This is not a full public website yet. It is a proof-of-feel prototype: a polished gateway screen, an immersive walkthrough mode, chapter controls, mobile-friendly navigation, and an Airbnb call to action. It is designed so Joey & Jo can continue project direction and approval, while Maaike & Laurens can review the House of Wander / Casa Cabane side: video, story, property details, permissions, and next content.",
    )

    doc.add_heading("Audience And Continuation Roles", level=1)
    for label, text in [
        ("Joey & Jo: ", "review the prototype direction, decide the next sprint scope, confirm the activation path, and decide when the project is ready for a wider team or public preview."),
        ("House of Wander team - Maaike & Laurens: ", "review the Casa Cabane experience from the property side, confirm video/source quality, approve factual details, provide missing content, and confirm what can be used beyond Airbnb."),
        ("Artificial Monks support: ", "keep the Dropbox project folder organized, integrate approved changes, run build/browser/mobile QA, and prepare clean handoffs for Claude, Codex, or future collaborators."),
    ]:
        bullet(doc, text, bold_label=label)

    doc.add_heading("Current Project Location", level=1)
    para(doc, "Canonical project folder:")
    code_line(doc, str(ROOT))
    para(doc, "Shortcut path:")
    code_line(doc, "/Users/wezienwel/houseofwander")
    para(
        doc,
        "The shortcut is only a symlink to the Dropbox folder. Future Codex and Claude work should use the Dropbox path as the real project home so all files stay together.",
    )

    doc.add_heading("What Is Built Now", level=1)
    for item in [
        "Cinematic Casa Cabane gateway screen with one primary action: Enter walkthrough.",
        "Full-screen guided walkthrough mode instead of a buried long-scroll section.",
        "Always-available controls: Exit, Previous, Play/Pause, Next, chapter filmstrip, timecode, and Airbnb CTA.",
        "Six chapters using the current walkthrough timing: Arrival, Exterior, Lake, Outdoor, Interior, Stay.",
        "Optional wheel/trackpad seeking inside the walkthrough while the page itself stays pinned.",
        "Exit returns the visitor to the top gateway state.",
        "Mobile-first layout for phone and tablet review.",
        "Reduced-motion behavior and video fallback state.",
    ]:
        bullet(doc, item)

    doc.add_heading("Naming And Copy Decisions", level=1)
    bullet(doc, "Current working brand spelling: House of Wander.", bold_label="")
    bullet(doc, "Current founder/reference spelling: Maaike & Laudi.", bold_label="")
    para(
        doc,
        "Older references to House of Wonder and Laudy were cleaned from active docs and app copy. Historical archive material may still contain older wording, but new work should use House of Wander and Maaike & Laudi unless the team decides otherwise.",
    )

    doc.add_heading("How To Run The Prototype", level=1)
    numbered(doc, "Open a terminal in the Dropbox project folder.")
    code_line(doc, f'cd "{ROOT}"')
    numbered(doc, "Install dependencies if needed.")
    code_line(doc, "npm install")
    numbered(doc, "Run a development preview.")
    code_line(doc, "npm run dev")
    numbered(doc, "Or run a clean production-style local preview.")
    code_line(doc, "npm run build")
    code_line(doc, "npm start -- -p 3000")
    para(doc, "If port 3000 is busy, use:")
    code_line(doc, "npm start -- -p 3001")
    para(doc, "Current verified preview during this work ran at:")
    code_line(doc, "http://localhost:3001")

    doc.add_heading("Files And Folders To Know", level=1)
    simple_table(
        doc,
        ["Path", "Purpose"],
        [
            ["app/page.tsx", "Main prototype interaction and page structure."],
            ["app/globals.css", "Visual design, responsive layout, and walkthrough overlay styling."],
            ["app/layout.tsx", "Page metadata."],
            ["public/media/casa-cabane-walkthrough.mp4", "Current higher-quality walkthrough video used by the prototype."],
            ["public/media/casa-cabane-poster.jpg", "Current poster image extracted from the updated video."],
            ["README.md", "Project overview and run commands."],
            ["CLAUDE.md", "Local instructions for Claude Code."],
            ["docs/claude-casa-cabane-handoff.md", "Prompt/handoff for Claude design iteration."],
            ["renders/casa-cabane-guided-qa/20260523-065751-v01/", "Browser QA screenshots from the guided prototype."],
        ],
        [2.55, 3.95],
    )

    doc.add_heading("Verification Completed", level=1)
    for item in [
        "npm run build passes with Next.js 16.2.6.",
        "Desktop QA completed in browser.",
        "Mobile QA completed at 390 x 844.",
        "Tablet QA completed at 768 x 1024.",
        "Enter walkthrough opens the full-screen journey.",
        "Previous, Next, and chapter buttons seek correctly.",
        "Exit returns to the top gateway state.",
        "No horizontal overflow was detected on tested mobile/tablet viewports.",
        "Reduced-motion mode remains usable and starts with Play instead of autoplay.",
        "Video fallback state is reachable and visible.",
        "Airbnb CTA keeps the existing Airbnb listing URL and opens in a new tab.",
    ]:
        bullet(doc, item)

    doc.add_heading("Current Limitations", level=1)
    for item in [
        "The current video is much higher quality, but it is also large at about 580 MB. Before public deployment, it may need a web-optimized export for faster loading.",
        "The page is local only. There is no public deployment, GitHub handoff, Vercel setup, analytics, CMS, or booking backend yet.",
        "The content is prototype copy, not final marketing or house-information copy.",
        "The walkthrough uses a single real video file. No Seedance/Higgsfield/Kie.ai loops are used in this version.",
        "The House of Wander collection landing page is not built yet.",
        "Amico/AI guidance is not implemented yet. That belongs to a later phase after the Casa Cabane experience direction is approved.",
    ]:
        bullet(doc, item)

    doc.add_heading("Better Casa Cabane Video: Current Status And Next Move", level=1)
    para(
        doc,
        "The better-quality Casa Cabane video has now been installed as public/media/casa-cabane-walkthrough.mp4. It was verified as 1080 x 1920 portrait, 60 fps, 254.4 seconds, H.264, around 580 MB. A new poster was extracted from the updated video and saved as public/media/casa-cabane-poster.jpg.",
    )
    for item in [
        "Keep the current installed video as the source for local review.",
        "Preserve the previous poster backup in the video-replacement QA render folder.",
        "QA the new file for orientation, crop, load time, poster match, chapter timing, and mobile caption readability.",
        "Before public deployment, create a smaller web-optimized MP4 if the 580 MB file feels too heavy.",
        "Reconfirm whether the installed file is the final approved source or still an intermediate export.",
    ]:
        numbered(doc, item)

    doc.add_heading("What The Team Needs To Decide", level=1)
    for item in [
        "Confirm the final public brand hierarchy: House of Wander, Casa Cabane, or another exact structure.",
        "Confirm whether Airbnb remains the only CTA for v1.",
        "Confirm whether the walkthrough video may be used internally only, externally, or publicly.",
        "Confirm whether Airbnb photos, review snippets, host names, and listing facts may be reused outside Airbnb.",
        "Confirm who approves final copy and visual direction: Joey & Jo, Maaike & Laurens, or a smaller approval group.",
        "Confirm whether the site should be English only, Dutch only, or bilingual.",
        "Confirm whether Casa Cabane should launch as standalone first or as the first property inside a wider House of Wander collection.",
        "Confirm whether direct inquiry or direct booking is a later goal.",
    ]:
        bullet(doc, item)

    doc.add_heading("Next Work Plan", level=1)
    phases = [
        ("Phase 1 - Casa Cabane Prototype Polish", "Optimize the current higher-quality video for web loading if needed, re-check chapter timings after any further media replacement or compression, keep the updated poster aligned, refine captions, and keep mobile-first QA as the main review standard."),
        ("Phase 2 - Full Casa Cabane Property Detail Layer", "Add practical stay details from Airbnb and host materials: capacity, amenities, rules, location, arrival/check-in info, and review proof."),
        ("Phase 3 - House of Wander Collection Landing", "Create the first collection landing page where visitors can see all accommodations and zoom into Casa Cabane as the first detailed prototype."),
        ("Phase 4 - Optional Amico Guidance", "Explore Amico as a light guided helper through properties and practical details, only after the Casa Cabane guided experience is approved."),
        ("Phase 5 - Public Preview / Deployment", "After local approval, connect GitHub and Vercel, rerun production QA, and decide whether the preview is private, passworded, or public."),
    ]
    for title, body in phases:
        doc.add_heading(title, level=2)
        para(doc, body)

    doc.add_heading("Recommended Agent Split", level=1)
    simple_table(
        doc,
        ["Claude", "Codex"],
        [
            ["High-taste design iteration.", "Integrating final code changes."],
            ["Copy refinement and spatial storytelling.", "Replacing and validating media assets."],
            ["Prompting for future mood loops or transition ideas.", "Running build checks and browser/mobile QA."],
            ["Exploring visual references from tutorials.", "Keeping the Dropbox project folder organized and preparing technical handoffs."],
        ],
        [3.15, 3.35],
    )

    doc.add_heading("Immediate To-Do List", level=1)
    simple_table(
        doc,
        ["Owner", "Next action"],
        [
            ["Joey & Jo", "Confirm final brand spelling and public naming structure."],
            ["Joey & Jo", "Decide whether the next build step is Casa Cabane detail content, the House of Wander collection landing, or public-preview preparation."],
            ["Joey & Jo", "Confirm whether Airbnb remains the v1 CTA."],
            ["House of Wander team - Maaike & Laurens", "Confirm whether the installed higher-quality walkthrough is the final approved source or whether an even cleaner export exists."],
            ["House of Wander team - Maaike & Laurens", "Confirm usage permission for video, photos, review snippets, host names, and Airbnb listing facts."],
            ["House of Wander team - Maaike & Laurens", "Provide or approve practical property details for the next Casa Cabane content layer."],
            ["Codex", "Rerun desktop, mobile, tablet, reduced-motion, and fallback QA after any further media replacement or compression."],
            ["Claude or design lead", "Refine the visual direction after the better video is in place."],
        ],
        [2.35, 4.15],
    )

    doc.add_heading("Current Recommendation", level=1)
    para(
        doc,
        "Do not expand into a full platform yet. Keep the next sprint focused on Casa Cabane, mobile-first walkthrough quality, video loading/performance, and the core guided feeling. Once that feels right, use Casa Cabane as the blueprint for the full House of Wander collection and the later Amico-guided experience.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

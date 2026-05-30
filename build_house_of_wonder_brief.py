from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "renders/house-of-wonder-brief/20260522-004801-v01/house-of-wonder-brief.docx"


SOURCES = {
    "airbnb_uk": "https://www.airbnb.co.uk/rooms/962896812968290131?unique_share_id=f38410d1-c5dc-4e02-b02b-d12ef25660cc&viralityEntryPoint=1&s=76",
    "airbnb_be": "https://www.airbnb.be/rooms/962896812968290131?guests=1&adults=1&s=67&unique_share_id=915e903d-cb45-4d1e-9776-0d9e07a52e61",
    "seedance": "https://seed.bytedance.com/en/seedance2_0",
    "seedance_launch": "https://seed.bytedance.com/en/blog/seedance-2-0-official-launch",
    "walkly": "https://www.walkly.app/en/products/examples",
    "tour_video": "https://tour.video/",
    "matterport": "https://matterport.com/en-us",
    "awwwards": "https://www.awwwards.com/awwwards/collections/storytelling/",
    "claude_code": "https://code.claude.com/docs/en/overview",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 20, "000000"),
        ("Heading 2", 16, "000000"),
        ("Heading 3", 14, "434343"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(18 if style_name != "Heading 1" else 20)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("House of Wander - Casa Cabane Stekene Project Activation Brief")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = False

    meta = doc.add_paragraph()
    meta.add_run("Prepared for: ").bold = True
    meta.add_run("Joey/internal project activation")
    meta.add_run(" | Date: ").bold = True
    meta.add_run("22 May 2026")
    meta.add_run(" | Status: ").bold = True
    meta.add_run("Kickoff brief with walkthrough video received")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        p.add_run(bold_label).bold = True
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            p.add_run(item[0]).bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            p.add_run(item[0]).bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = ""
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[idx], "F1F3F4")
        set_cell_border(hdr_cells[idx])
        hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(6)
    return table


def source_line(doc, label, source_key):
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    add_hyperlink(p, SOURCES[source_key], SOURCES[source_key])


def build():
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This document activates the first House of Wander prototype: an immersive, video-first property website for Casa Cabane in Stekene. The goal is to turn a standard rental listing into a guided digital experience where a visitor can feel the house, lake, garden, stillness, and practical stay details before clicking through to book or inquire.",
    )
    add_para(
        doc,
        "The brief is written for Joey and the internal team. It is designed to answer what we know publicly, what still needs to be researched, what the team must feed into the project, and how we should organize production before building the site.",
    )
    add_bullets(
        doc,
        [
            ("Primary outcome: ", "a clear activation plan that Joey can share and use to gather assets, align responsibilities, and start the project."),
            ("First property: ", "Casa Cabane - Stekene, described publicly on Airbnb as a stylish lakehouse in green nature."),
            ("Walkthrough status: ", "Maaike & Laudi's Casa Cabane walkthrough video has been received and can be used as the first prototype source asset."),
            ("Important boundary: ", "the available file appears to be a WhatsApp-exported version, so it is good for prototyping but the team should still provide the least-compressed original if possible."),
        ],
    )

    add_heading(doc, "2. Project Vision", 1)
    add_para(
        doc,
        "House of Wander should become a reusable visual platform for special stays. Instead of presenting Airbnbs as flat listings, the site should invite people into each place through video, atmosphere, room-by-room chapters, practical information, and a clean booking path.",
    )
    add_bullets(
        doc,
        [
            ("Experience principle: ", "show the feeling first, then make the practical booking information easy to find."),
            ("Prototype principle: ", "build one excellent property page before turning the system into a full multi-property platform."),
            ("Production principle: ", "keep the content model repeatable so later properties can be added without rebuilding the whole site."),
            ("Team principle: ", "use the codebase and briefing process as a shared learning track for Joey, Jo, and the hosts."),
        ],
    )

    add_heading(doc, "3. First Prototype: Casa Cabane - Stekene", 1)
    add_para(
        doc,
        "Public listing snapshot, accessed 22 May 2026. Airbnb content can change, so all operational details should be re-confirmed with Laudi, Maaike, Laurens, Thomas, or the live listing before launch.",
    )
    add_table(
        doc,
        ["Field", "Publicly visible information", "How it matters for the website"],
        [
            ["Public listing title", "Stylish lakehouse, green nature", "Sets the baseline SEO/title phrase; Casa Cabane can be the branded layer."],
            ["Location", "Stekene, Flemish Region, Belgium; described as between Antwerp and Ghent", "Use as access context and regional positioning."],
            ["Property type", "Entire home", "Supports a private-retreat story."],
            ["Capacity", "7 guests, 2 bedrooms, 5 beds, 1 bathroom", "Needed for practical specs and filters."],
            ["Comfort note", "Listing text says the house accommodates up to four adults comfortably", "Important to avoid overpromising in the new website copy."],
            ["Rating", "Guest favourite, 4.92 rating from 93 reviews", "Strong social proof for the first screen and trust section."],
            ["Host", "Hosted by Laurens, Superhost, 5 years hosting; co-hosts Thomas and Maaike", "Need team approval on which host names appear on House of Wander."],
            ["Signature features", "Lake, wild garden, outdoor tank bath, barbecue, fire bowl, outdoor dining, children's playhouse", "These become video chapters and visual hooks."],
            ["Core amenities", "Waterfront, kitchen, wifi, dedicated workspace, free parking", "Useful for quick-scan practical section."],
            ["Rules", "Check-in after 16:00, checkout before 11:00, 7 guests maximum", "Belongs in the practical stay details."],
        ],
        [1.45, 2.4, 2.45],
    )
    source_line(doc, "Source", "airbnb_uk")
    source_line(doc, "Dutch source mirror", "airbnb_be")

    add_heading(doc, "4. Received Walkthrough Video", 1)
    add_para(
        doc,
        "Maaike & Laudi's walkthrough video is now available locally and should become the base material for the first immersive walkthrough effect. A quick technical check confirms it is usable for planning and prototyping.",
    )
    add_table(
        doc,
        ["Field", "Current asset"],
        [
            ["File", "WhatsApp Video 2026-05-20 at 23.49.43 (1).mp4"],
            ["Location", "/Users/wezienwel/Downloads/WhatsApp Video 2026-05-20 at 23.49.43 (1).mp4"],
            ["Duration", "Approximately 4 minutes 14 seconds"],
            ["Size", "Approximately 43 MB"],
            ["Format", "MP4 container, H.264 video, AAC audio"],
            ["Resolution", "1024 x 576 encoded, with phone rotation metadata"],
            ["Visible walkthrough content", "Arrival/green lane, exterior approach, lake and wild garden, outdoor seating, glass facade, living area, kitchen/dining zone, interior lake views."],
        ],
        [1.65, 4.55],
    )
    add_heading(doc, "How to use this video in the prototype", 2)
    add_bullets(
        doc,
        [
            "Treat it as the first full walkthrough source, not just a reference clip.",
            "Cut it into chapters: arrival, exterior/lake, outdoor living, entrance, living room, kitchen/dining, views, final booking mood.",
            "Export poster stills from each chapter for navigation, loading states, and mobile fallbacks.",
            "Use the current WhatsApp file for prototype exploration, but request the original camera file if available for final launch quality.",
            "Check whether the audio should be kept, replaced with ambient sound/music, or muted under captions.",
            "Correct rotation and crop decisions before creating web-ready versions for desktop and mobile.",
        ],
    )

    add_heading(doc, "5. Public Research Snapshot", 1)
    add_para(
        doc,
        "Public research supports the direction, but it should not be copied blindly. The useful pattern is the combination of editorial storytelling, guided video, practical conversion, and repeatable development workflow.",
    )
    add_table(
        doc,
        ["Reference", "What it shows", "Use for House of Wander"],
        [
            ["Awwwards storytelling collection", "Strong storytelling sites merge visual design and UI to create interaction and engagement.", "Use as taste reference for pacing, transitions, and chapter-based narrative."],
            ["Walkly examples", "Interactive video walkthroughs can be embedded on websites and optimized for mobile viewers.", "Validate the idea that a property walkthrough can be more than a passive video."],
            ["Tour.video", "Guided video tours can pair video with CTAs, follow-ups, analytics, and lead conversion.", "Use as product reference for video plus action points, not just decoration."],
            ["Matterport", "Immersive 3D tours, floor plans, and property media help buyers or renters understand spaces.", "Consider whether Casa Cabane needs 3D/floor-plan later, but keep v1 video-first."],
            ["Seedance 2.0", "ByteDance positions it as multimodal video generation using text, image, audio, and video inputs with director-level control.", "Use only for concept tests, mood clips, or missing transition shots after rights are clear."],
            ["Claude Code docs", "Project instructions, MCP, skills, hooks, and repeatable workflows can make agent-assisted development more systematic.", "Translate this into repo instructions and workflow docs for the team."],
        ],
        [1.45, 2.4, 2.45],
    )

    add_heading(doc, "6. Experience Concept", 1)
    add_para(
        doc,
        "The Casa Cabane prototype should feel like arriving, entering, looking around, understanding the stay, and then deciding whether to book. The video is the main spatial anchor; text and UI should support it rather than compete with it.",
    )
    add_heading(doc, "Recommended first-screen flow", 2)
    add_numbered(
        doc,
        [
            ("Arrival mood: ", "quiet hero video or still frame of the lake/garden with the name Casa Cabane - Stekene."),
            ("Enter the house: ", "short scroll or click interaction that moves from exterior into interior."),
            ("Explore chapters: ", "living area, kitchen, bedroom, bathroom, garden/lake, outdoor bath, practical stay details."),
            ("Trust and proof: ", "ratings, review themes, host credibility, guest-favourite status."),
            ("Book or inquire: ", "clear CTA to Airbnb in v1 unless the team decides on a direct booking route."),
        ],
    )
    add_heading(doc, "Interaction ideas for v1", 2)
    add_bullets(
        doc,
        [
            "Video chapter navigation using room names and small poster stills.",
            "Scroll-timed captions that explain what the viewer is seeing without overloading the screen.",
            "A quiet practical panel that opens for beds, amenities, check-in, and house rules.",
            "Mobile-first vertical flow, with video compressed and captions readable on phone.",
            "Reduced-motion fallback for visitors who cannot or do not want immersive motion.",
        ],
    )

    add_heading(doc, "7. Content And Data We Need From The Team", 1)
    add_para(
        doc,
        "This is the most important activation checklist. The project can start with public research, but the website cannot become premium until the team provides real source assets and decisions.",
    )
    add_table(
        doc,
            ["Priority", "Needed input", "Owner to confirm", "Why we need it"],
            [
            ["P0", "Confirm permission to use Maaike & Laudi's received walkthrough video on the prototype", "Maaike / Laudi / Joey / hosts", "Main source for the immersive experience."],
            ["P0", "Original uncompressed walkthrough video if available", "Maaike / Laudi / Joey", "The WhatsApp version works for prototype, but final launch should use the cleanest possible source."],
            ["P0", "Permission to use the walkthrough video, Airbnb photos, host names, and guest-review snippets", "Hosts / Joey", "Avoid rights and privacy issues."],
            ["P0", "Booking decision: link to Airbnb only, direct inquiry, or both", "Jo / Joey / hosts", "Controls CTA, legal copy, tracking, and backend scope."],
            ["P0", "Final naming: House of Wander, Casa Cabane, or another brand hierarchy", "Jo / Joey", "Prevents inconsistent public naming."],
            ["P0", "High-resolution photos and video stills: exterior, lake, garden, living, kitchen, bedroom, bathroom, details", "Hosts / Joey", "Needed for hero, fallback states, posters, and social previews."],
            ["P0", "Accurate capacity and comfort wording", "Hosts", "Airbnb says 7 max but also notes up to four adults comfortably."],
            ["P1", "Brand assets: logo, fonts, colors, tone references, any existing domain/social handles", "Jo / Joey", "Defines the visual system and avoids placeholder branding."],
            ["P1", "Location story: nearby nature, borderlands, Antwerp/Ghent access, local tips", "Hosts", "Makes the site richer than an Airbnb copy-paste."],
            ["P1", "House rules, safety details, check-in flow, parking instructions, pet/smoking/children rules", "Hosts", "Needed for trust and practical usability."],
            ["P1", "Guest review highlights approved for use", "Hosts / Joey", "Public reviews are useful, but selected quotes need careful handling."],
            ["P1", "Access to a shared Google Drive folder for assets and approvals", "Joey", "Keeps production clean and lets Codex work from a single source of truth."],
            ["P2", "Domain, hosting, analytics, email, and GitHub/Vercel access", "Jo / Joey", "Needed once the team moves from prototype to real deployment."],
        ],
        [0.65, 2.25, 1.2, 2.2],
    )

    add_heading(doc, "8. Video And Media Production Needs", 1)
    add_para(
        doc,
        "The existing walkthrough can unlock the prototype, but the team should gather enough material for a polished edit and strong fallbacks.",
    )
    add_heading(doc, "Minimum video package", 2)
    add_bullets(
        doc,
        [
            "Received: one WhatsApp-exported Casa Cabane walkthrough video, suitable for prototype chaptering.",
            "Still needed if available: original camera export of the walkthrough, ideally 4K or 1080p, steady, with clean natural light.",
            "One vertical version or phone-safe crop for mobile and social snippets.",
            "Short exterior arrival shot: path, door, lake, or garden entry.",
            "Room passes: living, kitchen, bedroom, bathroom, garden/lake, outdoor bath.",
            "Detail shots: rain shower, bed texture, table, kitchen, fire bowl, water, trees, outdoor bath.",
            "Ambient audio if usable: birds, water, footsteps, door opening, quiet outside atmosphere.",
            "Poster stills for every chapter, exported as high-quality images.",
        ],
    )
    add_heading(doc, "Recommended shot rules", 2)
    add_bullets(
        doc,
        [
            "Film each room with a 3 to 5 second static opening before moving.",
            "Move slowly enough that the footage can be used behind text.",
            "Capture both practical truth and mood; do not hide spatial constraints.",
            "Avoid people in frame unless everyone has approved public use.",
            "Keep raw files, not only edited social exports.",
        ],
    )

    add_heading(doc, "9. Public Research We Still Need To Do", 1)
    add_para(
        doc,
        "Before design starts, we should expand the public research in a focused way. The goal is not to collect endless inspiration; it is to make decisions for this one prototype.",
    )
    add_numbered(
        doc,
        [
            ("Competitive property scan: ", "find 10 to 15 premium cabin, lakehouse, tiny-house, and boutique-stay websites in Belgium, Netherlands, France, and Scandinavia."),
            ("Interaction inspiration: ", "collect 8 to 12 visual websites that use video, scroll chapters, WebGL, or editorial room-by-room storytelling."),
            ("Booking pattern research: ", "compare Airbnb-only CTA, direct booking, inquiry form, and hybrid flows for small hospitality brands."),
            ("SEO and location research: ", "search phrases around Stekene, lakehouse Belgium, nature retreat, weekend getaway, Antwerp/Ghent escape, and Netherlands borderlands."),
            ("Rights-safe AI video research: ", "check whether Seedance or any AI-video tool can be used only for abstract mood tests or transitions without using protected faces, brands, or guest imagery."),
            ("Technical feasibility research: ", "test video compression, mobile performance, scroll interactions, and hosting costs before committing to heavy visuals."),
        ],
    )

    add_heading(doc, "10. Production Workflow", 1)
    add_para(
        doc,
        "The team should work from one shared source of truth. The cleanest setup is a Drive folder for assets and approvals, this Google Doc for the activation brief, and a repo that documents each decision as the prototype evolves.",
    )
    add_table(
        doc,
        ["Stage", "What happens", "Output"],
        [
            ["1. Intake", "Gather video, photos, listing facts, brand decisions, booking decision, access details.", "Approved asset folder and updated checklist."],
            ["2. Story map", "Convert assets into chapters: arrival, enter, explore, outside, trust, book.", "One-page content map and video chapter list."],
            ["3. Design prototype", "Create first visual direction and mobile/desktop flow.", "Prototype screens or live coded draft."],
            ["4. Codebase setup", "Set up Next.js/TypeScript project, content model, media pipeline, and project instructions.", "Working local repo Joey can follow."],
            ["5. First interactive build", "Build Casa Cabane page with real media and booking CTA.", "Internal preview URL."],
            ["6. QA and launch prep", "Test mobile, performance, links, copy, accessibility, analytics, and source attribution.", "Launch checklist and deployment decision."],
        ],
        [1.15, 3.05, 2.0],
    )

    add_heading(doc, "11. Technical Direction For The Prototype", 1)
    add_para(
        doc,
        "The first build should be ambitious visually but conservative structurally. Start with one property and a repeatable content schema, then scale once the Casa Cabane model proves itself.",
    )
    add_bullets(
        doc,
        [
            ("Recommended stack: ", "Next.js, TypeScript, Tailwind or a small design system, Framer Motion or GSAP only where interactions need it."),
            ("Content model: ", "structured property data plus markdown/JSON for sections, chapters, amenities, CTAs, reviews, and source links."),
            ("Media pipeline: ", "video compression, poster images, captions, lazy loading, mobile variants, and reduced-motion fallbacks."),
            ("CMS decision: ", "start with simple repo-managed content for v1; move to Notion, Sanity, or another CMS only when multiple properties are coming."),
            ("Backend needs: ", "v1 may only need tracking and external booking link; direct booking, inquiry forms, CRM, payments, or availability sync would be a separate phase."),
            ("Performance rule: ", "the site must feel premium on phone, not only on a fast laptop."),
        ],
    )

    add_heading(doc, "12. Open Questions To Resolve Before Build", 1)
    add_bullets(
        doc,
        [
            "Is the public brand House of Wander, Casa Cabane, or another exact spelling?",
            "Should the first public CTA send visitors to Airbnb, or should the site collect direct inquiries?",
            "Who approves copy: Jo, Joey, Laudi, Maaike, Laurens, Thomas, or a smaller group?",
            "Can we use selected guest review snippets outside Airbnb, and under what wording?",
            "What is the exact ownership and usage permission for the walkthrough video and Airbnb photos?",
            "Can Maaike & Laudi provide the original non-WhatsApp-compressed video file?",
            "Do we need the site in English only, Dutch only, or bilingual?",
            "Will Casa Cabane be presented as part of a future collection, or as a standalone first launch?",
            "What is the deadline for first clickable prototype, internal review, and public launch?",
        ],
    )

    add_heading(doc, "13. Immediate Next Actions", 1)
    add_numbered(
        doc,
        [
            "Joey shares this brief with Jo and the property team.",
            "Create a shared Google Drive folder named House of Wander - Casa Cabane - Assets.",
            "Upload the received walkthrough video, then add the original uncompressed file if Maaike & Laudi still have it.",
            "Confirm the exact brand spelling and booking route.",
            "Assign one person to verify public Airbnb facts against the hosts.",
            "Codex turns the approved assets into a story map, content schema, and first implementation plan.",
            "Team reviews the prototype scope before any public build or deployment.",
        ],
    )

    add_heading(doc, "14. Source Links Used In This Brief", 1)
    for label, key in [
        ("Airbnb UK listing", "airbnb_uk"),
        ("Airbnb Belgium listing", "airbnb_be"),
        ("ByteDance Seedance 2.0 model page", "seedance"),
        ("ByteDance Seedance 2.0 official launch blog", "seedance_launch"),
        ("Walkly interactive video walkthrough examples", "walkly"),
        ("Tour.video virtual leasing/video tour reference", "tour_video"),
        ("Matterport immersive 3D property marketing reference", "matterport"),
        ("Awwwards storytelling collection", "awwwards"),
        ("Claude Code overview for workflow inspiration", "claude_code"),
    ]:
        source_line(doc, label, key)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Appendix: Intake Message Joey Can Send", 1)
    add_para(
        doc,
        "Hi all, to activate the House of Wander / Casa Cabane prototype, we now have Maaike & Laudi's walkthrough video as the first source asset. Please upload it to the shared project folder, and if possible also add the original non-WhatsApp-compressed version. We also need high-resolution photos, any existing brand material, approved listing text, house rules, host names, booking preference, and any local tips or story notes. The goal is to turn Casa Cabane into a first immersive video-first property website prototype, not just copy the Airbnb page.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
